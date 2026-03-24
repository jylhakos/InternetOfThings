import os
import uuid
import aiofiles
from typing import List, Dict, Any, Optional, BinaryIO
from pathlib import Path
from datetime import datetime
import hashlib
from loguru import logger

# Document processing libraries
import PyPDF2
from docx import Document as DocxDocument
import markdown
from io import BytesIO

from app.core.config import settings
from app.models.schemas import (
    Document, DocumentChunk, DocumentType, 
    DocumentMetadata, DocumentResponse
)
from app.services.vector_service import VectorService


class DocumentService:
    """Service for document processing and management"""
    
    def __init__(self, vector_service: VectorService):
        self.vector_service = vector_service
        self.upload_dir = Path("uploads")
        self.upload_dir.mkdir(exist_ok=True)
    
    def _get_file_type(self, filename: str) -> Optional[DocumentType]:
        """Determine file type from filename"""
        extension = Path(filename).suffix.lower()
        
        type_mapping = {
            '.pdf': DocumentType.PDF,
            '.txt': DocumentType.TEXT,
            '.docx': DocumentType.DOCX,
            '.md': DocumentType.MARKDOWN,
            '.markdown': DocumentType.MARKDOWN
        }
        
        return type_mapping.get(extension)
    
    def _validate_file(self, filename: str, file_size: int, content_type: str) -> bool:
        """Validate uploaded file"""
        # Check file size
        if file_size > settings.MAX_FILE_SIZE:
            raise ValueError(f"File too large. Maximum size: {settings.MAX_FILE_SIZE} bytes")
        
        # Check file type
        if content_type not in settings.ALLOWED_FILE_TYPES:
            raise ValueError(f"Unsupported file type: {content_type}")
        
        # Check filename
        if not filename or len(filename) > 255:
            raise ValueError("Invalid filename")
        
        return True
    
    async def extract_text_from_pdf(self, file_content: bytes) -> str:
        """Extract text from PDF file"""
        try:
            pdf_reader = PyPDF2.PdfReader(BytesIO(file_content))
            text = ""
            
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            
            return text.strip()
        except Exception as e:
            logger.error(f"Failed to extract text from PDF: {e}")
            raise ValueError("Unable to process PDF file")
    
    async def extract_text_from_docx(self, file_content: bytes) -> str:
        """Extract text from DOCX file"""
        try:
            doc = DocxDocument(BytesIO(file_content))
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            return text.strip()
        except Exception as e:
            logger.error(f"Failed to extract text from DOCX: {e}")
            raise ValueError("Unable to process DOCX file")
    
    async def extract_text_from_markdown(self, file_content: bytes) -> str:
        """Extract text from Markdown file"""
        try:
            md_content = file_content.decode('utf-8')
            # Convert markdown to plain text (removing markup)
            html = markdown.markdown(md_content)
            # Simple HTML tag removal (you might want to use BeautifulSoup for better results)
            import re
            text = re.sub(r'<[^>]+>', '', html)
            return text.strip()
        except Exception as e:
            logger.error(f"Failed to extract text from Markdown: {e}")
            raise ValueError("Unable to process Markdown file")
    
    async def extract_text(self, file_content: bytes, file_type: DocumentType) -> str:
        """Extract text content based on file type"""
        if file_type == DocumentType.PDF:
            return await self.extract_text_from_pdf(file_content)
        elif file_type == DocumentType.DOCX:
            return await self.extract_text_from_docx(file_content)
        elif file_type == DocumentType.MARKDOWN:
            return await self.extract_text_from_markdown(file_content)
        elif file_type == DocumentType.TEXT:
            return file_content.decode('utf-8')
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
    
    def chunk_text(self, text: str, chunk_size: int = None, overlap: int = None) -> List[str]:
        """Split text into chunks for vector storage"""
        chunk_size = chunk_size or settings.MAX_CHUNK_SIZE
        overlap = overlap or settings.CHUNK_OVERLAP
        
        if len(text) <= chunk_size:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            # Find the end of the chunk
            end = start + chunk_size
            
            # If we're not at the end of the text, try to break at a sentence or word boundary
            if end < len(text):
                # Look for sentence ending
                sentence_end = text.rfind('.', start, end)
                if sentence_end > start + chunk_size // 2:
                    end = sentence_end + 1
                else:
                    # Look for word boundary
                    word_end = text.rfind(' ', start, end)
                    if word_end > start + chunk_size // 2:
                        end = word_end
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            # Move start position with overlap
            start = end - overlap if end - overlap > start else end
        
        return chunks
    
    async def save_file(self, filename: str, file_content: bytes) -> str:
        """Save uploaded file to disk"""
        # Generate unique filename
        file_id = str(uuid.uuid4())
        file_extension = Path(filename).suffix
        unique_filename = f"{file_id}{file_extension}"
        file_path = self.upload_dir / unique_filename
        
        async with aiofiles.open(file_path, 'wb') as f:
            await f.write(file_content)
        
        return unique_filename
    
    async def process_document(
        self,
        filename: str,
        file_content: bytes,
        content_type: str,
        metadata: Optional[DocumentMetadata] = None
    ) -> DocumentResponse:
        """Process and store document"""
        try:
            # Validate file
            file_size = len(file_content)
            self._validate_file(filename, file_size, content_type)
            
            # Determine file type
            file_type = self._get_file_type(filename)
            if not file_type:
                raise ValueError("Unable to determine file type")
            
            # Extract text content
            text_content = await self.extract_text(file_content, file_type)
            
            if not text_content.strip():
                raise ValueError("No text content found in file")
            
            # Create chunks
            chunks = self.chunk_text(text_content)
            
            # Generate document ID
            doc_id = str(uuid.uuid4())
            
            # Save file to disk
            saved_filename = await self.save_file(filename, file_content)
            
            # Prepare metadata
            if not metadata:
                metadata = DocumentMetadata()
            
            if not metadata.title:
                metadata.title = Path(filename).stem
            
            # Create document chunks for vector storage
            document_chunks = []
            for i, chunk_text in enumerate(chunks):
                chunk_metadata = {
                    "document_id": doc_id,
                    "filename": filename,
                    "chunk_index": i,
                    "total_chunks": len(chunks),
                    "file_type": file_type.value,
                    **metadata.dict()
                }
                
                chunk = DocumentChunk(
                    id=f"{doc_id}_chunk_{i}",
                    content=chunk_text,
                    metadata=chunk_metadata
                )
                document_chunks.append(chunk)
            
            # Store chunks in vector database
            success = await self.vector_service.add_documents(document_chunks)
            if not success:
                raise Exception("Failed to store document in vector database")
            
            # Create document object
            document = Document(
                id=doc_id,
                filename=saved_filename,
                file_type=file_type,
                file_size=file_size,
                content_preview=text_content[:500] + "..." if len(text_content) > 500 else text_content,
                metadata=metadata,
                chunk_count=len(chunks),
                created_at=datetime.now(),
                updated_at=datetime.now()
            )
            
            logger.info(f"Successfully processed document: {filename} ({len(chunks)} chunks)")
            
            return DocumentResponse(
                document=document,
                message="Document processed and stored successfully"
            )
            
        except Exception as e:
            logger.error(f"Failed to process document {filename}: {e}")
            raise Exception(f"Document processing failed: {str(e)}")
    
    async def delete_document(self, document_id: str) -> bool:
        """Delete document and its chunks"""
        try:
            # Delete from vector database (this will delete all chunks with this document_id)
            # Note: You might need to implement a method to delete by metadata filter
            success = await self.vector_service.delete_document(document_id)
            
            if success:
                logger.info(f"Successfully deleted document: {document_id}")
            
            return success
            
        except Exception as e:
            logger.error(f"Failed to delete document {document_id}: {e}")
            return False
    
    def calculate_file_hash(self, file_content: bytes) -> str:
        """Calculate SHA-256 hash of file content"""
        return hashlib.sha256(file_content).hexdigest()
    
    async def get_document_info(self, filename: str) -> Dict[str, Any]:
        """Get document information from filename"""
        file_path = self.upload_dir / filename
        
        if not file_path.exists():
            return {}
        
        stat = file_path.stat()
        
        return {
            "filename": filename,
            "size": stat.st_size,
            "created": datetime.fromtimestamp(stat.st_ctime),
            "modified": datetime.fromtimestamp(stat.st_mtime),
            "exists": True
        }
    
    async def list_uploaded_files(self) -> List[Dict[str, Any]]:
        """List all uploaded files"""
        files = []
        
        for file_path in self.upload_dir.iterdir():
            if file_path.is_file():
                stat = file_path.stat()
                files.append({
                    "filename": file_path.name,
                    "size": stat.st_size,
                    "created": datetime.fromtimestamp(stat.st_ctime),
                    "modified": datetime.fromtimestamp(stat.st_mtime),
                })
        
        return files
    
    async def cleanup_orphaned_files(self) -> int:
        """Remove files that are not referenced in vector database"""
        # This would require implementing a way to track which files are in use
        # For now, we'll just return 0
        return 0
